import os
import shutil
import fnmatch
import numpy as np
import pandas as pd
import string
from urllib.request import urlretrieve
from docx import Document
from docx.shared import Inches


name_csv='test.csv'
main_path = name_csv.split('.')
main_path = main_path[0]
data=pd.read_csv(name_csv)
#print(data)
path = r"/Users/ivoyoung/Documents/GitHub/Recruitment_2020"

def word_doc_creator(row):

    document = Document()

    document.add_heading(str(row[3]), 0)


    #Form_id                0
    #Form_date              1
    #Status                 2
    #Name                   3
    #Email                  4
    #Phone-number           5
    #Matriculation-number   6
    #Degree-subject         7
    #Study-year             8
    #Degree-type            9
    #Gpa                    10
    #Option-1               11
    #Option-2               12
    #Experience             13
    #Why-join               14
    #Helpful-application    15
    #Interview-times        16
    #Hear-about             17
    #Applied-before         18


    p = document.add_paragraph(str(row[14]))







    doc_name=str(row[3])+'.docx'
    document.save(doc_name)
    old_path1=path+'/'+doc_name
    new_path1=dir_path+'/'+doc_name


    shutil.move(old_path1,new_path1)

    return 0


def txt_doc_creator(row):


    f  = open(str(row[3])+ " "+str(row[4])+".txt", "w+")

    f.write(("Form_id: "+str(row[0])))
    f.write("\n")
    f.write(("Form_date: "+str(row[1])))
    f.write("\n")
    f.write(("Status: "+str(row[2])))
    f.write("\n")
    f.write(("Name: "+str(row[3])))
    f.write("\n")
    f.write(("Email: "+str(row[4])))
    f.write("\n")
    f.write(("Phone-number: "+str(row[5])))
    f.write("\n")
    f.write(("Matriculation-number: "+str(row[6])))
    f.write("\n")
    f.write(("Degree-subject: "+str(row[7])))
    f.write("\n")
    f.write(("Degree-type: "+str(row[8])))
    f.write("\n")
    f.write(("Study-year: "+str(row[9])))
    f.write("\n")
    f.write(("GPA: "+str(row[10])))
    f.write("\n")
    f.write(("Option-1:  "+str(row[11])))
    f.write("\n")
    f.write(("Option-2: "+str(row[12])))
    f.write("\n")
    f.write(("Experience:  "+str(row[13])))
    f.write("\n")
    f.write(("Why-join:  "+str(row[14])))
    f.write("\n")
    f.write(("Helpful-application:  "+str(row[15])))
    f.write("\n")
    f.write(("Interview-times:  "+str(row[16])))
    f.write("\n")
    f.write(("Hear-about:  "+str(row[17])))
    f.write("\n")
    f.write(("Applied-before:  "+str(row[18])))

    
    #Phone-number           5
    #   6
    #         7
    #Study-year             8
    #Degree-type            9
    #Gpa                    10
    #Option-1               11
    #Option-2               12
    #Experience             13
    #Why-join               14
    #Helpful-application    15
    #Interview-times        16
    #Hear-about             17
    #Applied-before         18
    
    
    
    f.close()




    old_path=path+'/'+str(row[3])+ " "+str(row[4])+".txt"
    new_path=dir_path+'/'+str(row[3])+ " "+str(row[4])+".txt"
    shutil.move(old_path,new_path)

    return 0

for index, row in data.iterrows():
    new_dir=str(row[3]) + " "+str(row[4])
    dir_path = "{0:s}/{1:s}".format(path, new_dir)
    try:
        os.mkdir(dir_path)
    except FileExistsError as err:
        print("Folder Exists")


    txt_doc_creator(row)
    word_doc_creator(row)
    download_url=str(row[19])
    urlretrieve(download_url, dir_path+'/'+str(row[3])+'.pdf')


    


#Form_id                0
#Form_date              1
#Status                 2
#Name                   3
#Email                  4
#Phone-number           5
#Matriculation-number   6
#Degree-subject         7
#Study-year             8
#Degree-type            9
#Gpa                    10
#Option-1               11
#Option-2               12
#Experience             13
#Why-join               14
#Helpful-application    15
#Interview-times        16
#Hear-about             17
#Applied-before         18
#Cv                     19

#if not os.path.exists(main_path):
#    os.makedirs(main_path)

#for index, row in data.iterrows():
#
#    f  = open(str(row[3])+ " "+str(row[4])+".txt", "w+")

#    f.write(("Form_id: "+str(row[0])+ " Form_date:"+str(row[1])))
#    f.write("\n")
#    f.write("\n")
#    f.write("\n")
#    f.write(("Status: "+str(row[2])+ " Name:"+str(row[3])))
#    f.write("\n")
#    f.write(("Email: "+str(row[4])+ " Phone-number:"+str(row[5])))
#    f.close()

    
    #for i in range(10):
    #    f.write("This is line %d\r\n" % (i+1))
    #file_path=main_path+"/" + str(row[3])
    #print(file_path)
    
    #try:
    #    os.makedirs(main_path+"/" + str(row[3] + " "+str(row[4])))
        
    #    os.chdir(main_path+"/" + str(row[3] + " "+str(row[4])))
    #    with open(str(row[3] + " "+str(row[4])+'.txt'),'w'):
    #        print('yes')

    #except:
    #    pass
    #except:
    #    pass

    #except:
    #    pass
    #print(row)
