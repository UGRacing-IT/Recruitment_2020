import os
import shutil
import fnmatch
import numpy as np
import pandas as pd
import string
from urllib.request import urlretrieve
from docx import Document
from docx.shared import Inches


#Gets the working directory 
path = os.getcwd()

#Finds the csv downloaded from the website, will get only one csv in the current directory
for file in os.listdir(path):
    if fnmatch.fnmatch(file, '*.csv'):
            print(file)
            data=pd.read_csv(file)


def word_doc_creator(row):

    #takes a row from the csv and creates a word doc with all the Relevant info

    document = Document()

    document.add_heading(str(row[3]), 0)


    document.add_paragraph(str(row[4])+'                                                              '+str(row[5]))

    items = (
    ('Form_id:', str(row[0])),
    ('Form_date:', str(row[1])),
    ('Matriculation-number:',str(row[6])),
    ('Degree-subject:',str(row[7])),
    ('Study-year:',str(row[8])),
    ('Degree-type:',str(row[9])),
    ('GPA:',str(row[10])),
    ('Option-1:',str(row[11])),
    ('Option-2:',str(row[12])),
    ('Hear-about:',str(row[17])),
    ('Applied-before:',str(row[18]))
    )

    table = document.add_table(1, 2)

    # populate header row --------
    #heading_cells = table.rows[0].cells
    #heading_cells[0].text = 'Qty'
    #heading_cells[1].text = 'SKU'

    # add a data row for each item
    for item in items:
        cells = table.add_row().cells
        cells[0].text = str(item[0])
        cells[1].text = str(item[1])

    document.add_heading('Experience', level=2)
    p = document.add_paragraph(str(row[13]))

    document.add_heading('Why-join', level=2)
    p = document.add_paragraph(str(row[14]))

    document.add_heading('Helpful-application', level=2)
    p = document.add_paragraph(str(row[15]))

    document.add_heading('Interview-times', level=2)
    p = document.add_paragraph(str(row[16]))




    doc_name=str(row[3])+'.docx'
    document.save(doc_name)
    old_path1=path+'/'+doc_name
    new_path1=dir_path+'/'+doc_name


    shutil.move(old_path1,new_path1)

    return 0


def txt_doc_creator(row):


    f  = open(str(row[3])+ " "+str(row[4])+".txt", "w+")

    f.write(("Form_id: "+str(row[0])))
    f.write("\n")
    f.write(("Form_date: "+str(row[1])))
    f.write("\n")
    f.write(("Status: "+str(row[2])))
    f.write("\n")
    f.write(("Name: "+str(row[3])))
    f.write("\n")
    f.write(("Email: "+str(row[4])))
    f.write("\n")
    f.write(("Phone-number: "+str(row[5])))
    f.write("\n")
    f.write(("Matriculation-number: "+str(row[6])))
    f.write("\n")
    f.write(("Degree-subject: "+str(row[7])))
    f.write("\n")
    f.write(("Study-year: "+str(row[8])))
    f.write("\n")
    f.write(("Degree-type: "+str(row[9])))
    f.write("\n")
    f.write(("GPA: "+str(row[10])))
    f.write("\n")
    f.write(("Option-1:  "+str(row[11])))
    f.write("\n")
    f.write(("Option-2: "+str(row[12])))
    f.write("\n")
    f.write(("Experience:  "+str(row[13])))
    f.write("\n")
    f.write(("Why-join:  "+str(row[14])))
    f.write("\n")
    f.write(("Helpful-application:  "+str(row[15])))
    f.write("\n")
    f.write(("Hear-about:  "+str(row[16])))
    f.write("\n")
    f.write(("Applied-before:  "+str(row[17])))  
    f.close()
    
    #Moves the txt file to the relevant folder
    old_path=path+'/'+str(row[3])+ " "+str(row[4])+".txt"
    new_path=dir_path+'/'+str(row[3])+ " "+str(row[4])+".txt"
    shutil.move(old_path,new_path)

    return 0

for index, row in data.iterrows():
    #Iterates over all the rows in the csv
    #Creates a new folder
    #if multiple submissions are made with the 
    #same name and email it is assumed that
    #the same person is Applying again
    #and so the folder already exsists 
    #and the most recent submission is
    #the one that it takes
    new_dir=str(row[3]) + " "+str(row[4]+" "+str(row[11]))
    dir_path = "{0:s}/{1:s}".format(path, new_dir)
    try:
        os.mkdir(dir_path)
    except FileExistsError as err:
        print("Folder Exists")


    txt_doc_creator(row)
    word_doc_creator(row)
    download_url=str(row[18])
    #downloads the uploaded csv
    print(row[3])
    try:
        #This wont work on eduroam HTTP Error 403
        urlretrieve(download_url, dir_path+'/'+str(row[3])+'.pdf')
    except:
        #'deals' with unicode error
        print('Error downloading: ',download_url)
        pass

